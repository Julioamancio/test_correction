import os
from flask import Flask, request, redirect, url_for, render_template, session, send_file, flash
from werkzeug.utils import secure_filename
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document
from io import BytesIO

app = Flask(__name__)
app.secret_key = 'change-me'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///app.db'
app.config['UPLOAD_FOLDER'] = 'uploads'

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)

class Professor(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(128), nullable=False)

class Question(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    professor_id = db.Column(db.Integer, db.ForeignKey('professor.id'), nullable=False)
    text = db.Column(db.Text, nullable=False)
    option_a = db.Column(db.Text, nullable=False)
    option_b = db.Column(db.Text, nullable=False)
    option_c = db.Column(db.Text, nullable=False)
    option_d = db.Column(db.Text, nullable=False)
    option_e = db.Column(db.Text, nullable=False)
    correct_option = db.Column(db.String(1), nullable=False)
    image_filename = db.Column(db.String(200))
    classification = db.Column(db.String(100))

    professor = db.relationship('Professor')

def init_db():
    db.create_all()

@app.route('/')
def index():
    if 'professor_id' in session:
        return redirect(url_for('list_questions'))
    return redirect(url_for('login'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        if Professor.query.filter_by(username=username).first():
            flash('User exists')
            return render_template('register.html')
        prof = Professor(username=username, password_hash=generate_password_hash(password))
        db.session.add(prof)
        db.session.commit()
        flash('Registered successfully')
        return redirect(url_for('login'))
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        user = Professor.query.filter_by(username=request.form['username']).first()
        if user and check_password_hash(user.password_hash, request.form['password']):
            session['professor_id'] = user.id
            return redirect(url_for('list_questions'))
        flash('Invalid credentials')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('professor_id', None)
    return redirect(url_for('login'))

@app.route('/questions')
def list_questions():
    if 'professor_id' not in session:
        return redirect(url_for('login'))
    prof_id = session['professor_id']
    qs = Question.query.filter_by(professor_id=prof_id).all()
    return render_template('list_questions.html', qs=qs)

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_file(os.path.join(app.config['UPLOAD_FOLDER'], filename))

@app.route('/questions/new', methods=['GET', 'POST'])
def new_question():
    if 'professor_id' not in session:
        return redirect(url_for('login'))
    if request.method == 'POST':
        img = request.files.get('image')
        filename = None
        if img:
            filename = secure_filename(img.filename)
            img.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        q = Question(
            professor_id=session['professor_id'],
            text=request.form['text'],
            option_a=request.form['option_a'],
            option_b=request.form['option_b'],
            option_c=request.form['option_c'],
            option_d=request.form['option_d'],
            option_e=request.form['option_e'],
            correct_option=request.form['correct'],
            image_filename=filename,
            classification=request.form.get('classification')
        )
        db.session.add(q)
        db.session.commit()
        return redirect(url_for('list_questions'))
    return render_template('new_question.html')

@app.route('/questions/import', methods=['GET', 'POST'])
def import_questions():
    if 'professor_id' not in session:
        return redirect(url_for('login'))
    if request.method == 'POST':
        docx_file = request.files['docx']
        if docx_file:
            document = Document(docx_file)
            img_count = 0
            for para in document.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                # Very simplistic parser: expect paragraphs in order
                if text.startswith('Q:'):
                    q_text = text[2:].strip()
                elif text.startswith('A:'):
                    opt_a = text[2:].strip()
                elif text.startswith('B:'):
                    opt_b = text[2:].strip()
                elif text.startswith('C:'):
                    opt_c = text[2:].strip()
                elif text.startswith('D:'):
                    opt_d = text[2:].strip()
                elif text.startswith('E:'):
                    opt_e = text[2:].strip()
                elif text.startswith('Correct:'):
                    correct = text.split(':')[1].strip()
                    q = Question(
                        professor_id=session['professor_id'],
                        text=q_text,
                        option_a=opt_a,
                        option_b=opt_b,
                        option_c=opt_c,
                        option_d=opt_d,
                        option_e=opt_e,
                        correct_option=correct,
                        image_filename=None
                    )
                    db.session.add(q)
            db.session.commit()
            flash('Questions imported')
        return redirect(url_for('list_questions'))
    return render_template('import_questions.html')

@app.route('/questions/export')
def export_questions():
    if 'professor_id' not in session:
        return redirect(url_for('login'))
    ids = request.args.get('ids', '')
    id_list = [int(i) for i in ids.split(',') if i.isdigit()]
    qs = Question.query.filter(Question.professor_id == session['professor_id'], Question.id.in_(id_list)).all()
    doc = Document()
    for q in qs:
        doc.add_paragraph(q.text)
        if q.image_filename:
            doc.add_picture(os.path.join(app.config['UPLOAD_FOLDER'], q.image_filename))
        doc.add_paragraph(f"A: {q.option_a}")
        doc.add_paragraph(f"B: {q.option_b}")
        doc.add_paragraph(f"C: {q.option_c}")
        doc.add_paragraph(f"D: {q.option_d}")
        doc.add_paragraph(f"E: {q.option_e}")
        doc.add_paragraph(f"Correct: {q.correct_option}")
        doc.add_paragraph('---')
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name='questions.docx')

if __name__ == '__main__':
    with app.app_context():
        init_db()
    app.run(debug=True)
